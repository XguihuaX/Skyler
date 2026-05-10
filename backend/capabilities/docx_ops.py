"""v3.5 chunk 7 — docx 操作 capability（姿态 A demo：本地 capability + SAFE 沙箱）。

3 个 capability：

* ``docx.create``  —— 创建新 .docx（title + paragraphs）
* ``docx.read``    —— 读 .docx 内容（title 启发式 + 段落列表 + 词数）
* ``docx.append`` —— 末尾追加段落

设计与现有 capability 完全对齐：

* ``@register_capability`` 装饰器（注册到 ``CapabilityRegistry`` + ``ToolRegistry``）
* handler ``async def(**_kwargs)`` —— 兜住 ChatAgent 注入的 ``user_id``
* description 走 chunk 1.7 verbatim 引导风格（强引导 + 触发规则）
* 所有错误返回 ``{"error": "<code>", ...}`` dict，不 raise（LLM 自己根据
  上下文回话；与 ``backend/capabilities/clipboard.py`` 同模式）

### SAFE 沙箱

默认 ``~/Documents/Skyler/docs/``（用户可见，方便手动取文件）；
``config.yaml`` ``skills.docx.safe_dir`` 可覆盖。首次启动自动创建 + ``mode=0o700``
（macOS / Linux），通过 ``backend/utils/safe_path.py`` 集中实现 path traversal
防御。

### 不支持

* 图片 / 表格 / 公式 / 复杂排版（chunk 7 v1 spec：纯文本段落 + 一级标题）
* 子目录（filename 只能是 stem，``safe_resolve(allow_subdirs=False)``）
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from backend.capabilities import Consumer, TriggerMode, register_capability
from backend.config import config_yaml
from backend.utils.safe_path import ensure_sandbox_dir, safe_resolve

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Sandbox resolution
# ---------------------------------------------------------------------------

def _resolve_safe_dir() -> Path:
    """读 ``config.yaml`` skills.docx.safe_dir，缺省 ``~/Documents/Skyler/docs/``。

    每次调用都重新读 config（支持热重载），但 ``ensure_sandbox_dir`` 幂等。
    """
    skills_cfg = (config_yaml.get("skills") or {}).get("docx") or {}
    raw = skills_cfg.get("safe_dir") or "~/Documents/Skyler/docs"
    return Path(str(raw)).expanduser()


def _safe_dir() -> Path:
    base = _resolve_safe_dir()
    return ensure_sandbox_dir(base, mode=0o700)


def _normalize_filename(filename: str) -> str:
    """若 LLM 没带 ``.docx`` 后缀，自动补；其他后缀拒。

    例：``"周报_2026年05月"`` → ``"周报_2026年05月.docx"``。
    """
    s = filename.strip()
    lower = s.lower()
    if lower.endswith(".docx"):
        return s
    if "." in Path(s).name:
        # 用户带了别的后缀（``.doc`` / ``.txt``）→ 拒
        raise ValueError("filename must end with .docx (or no extension)")
    return s + ".docx"


# ---------------------------------------------------------------------------
# 1. docx.create
# ---------------------------------------------------------------------------

@register_capability(
    name="docx.create",
    display_name="创建 Word 文档",
    description=(
        "创建一份新的 Word 文档（.docx），保存到 Skyler 文档沙箱目录。"
        "适用场景：用户说「帮我写一份…」「起草一个文档」「做个周报」。\n\n"
        "参数：\n"
        "- filename: 文件名（你按内容自己起名，如 ``周报_2026年05月.docx``。"
        "可不带后缀，会自动补 .docx；不能含路径分隔符）\n"
        "- title: 文档一级大标题（一句话，不能为空）\n"
        "- paragraphs: 正文段落列表（list[str]，每段一项；可为空表示只要标题）\n\n"
        "返回 ``{path, size_bytes}``。若 filename 重复会**覆盖原文件**——需要"
        "保留旧版本时让用户先确认。"
    ),
    category="files",
    consumers=[Consumer.CHAT_AGENT],
    trigger_modes=[TriggerMode.ON_DEMAND],
    icon="file-text",
    parameters_schema={
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "Word 文件名（自动补 .docx 后缀），不能含路径分隔符",
            },
            "title": {
                "type": "string",
                "description": "文档一级大标题",
            },
            "paragraphs": {
                "type": "array",
                "items": {"type": "string"},
                "description": "正文段落列表（可空）",
            },
        },
        "required": ["filename", "title"],
    },
)
async def docx_create(
    filename: str = "",
    title: str = "",
    paragraphs: list[str] | None = None,
    **_kwargs: Any,
) -> dict:
    if not filename or not filename.strip():
        return {"error": "missing_filename"}
    if not title or not title.strip():
        return {"error": "missing_title"}
    try:
        canonical = _normalize_filename(filename)
        target = safe_resolve(_safe_dir(), canonical, allow_subdirs=False)
    except ValueError as exc:
        return {"error": "invalid_path", "detail": str(exc)}

    paragraphs = paragraphs or []
    try:
        doc = Document()
        doc.add_heading(title.strip(), level=1)
        for p in paragraphs:
            doc.add_paragraph(str(p))
        doc.save(str(target))
    except Exception as exc:
        logger.exception("[docx.create] save failed for %s", target)
        return {"error": "save_failed", "detail": str(exc)}

    size = target.stat().st_size
    rel = target.name
    logger.info("[docx.create] wrote %s (%d bytes)", target, size)
    return {"path": rel, "size_bytes": size}


# ---------------------------------------------------------------------------
# 2. docx.read
# ---------------------------------------------------------------------------

@register_capability(
    name="docx.read",
    display_name="读取 Word 文档",
    description=(
        "读取沙箱中已有的 Word 文档内容。适用场景：用户说「读一下我的XX文档」"
        "「看看那个周报里都写了啥」「总结一下 XX 文档」。\n\n"
        "参数：\n"
        "- filename: 文件名（带不带 .docx 都行）\n\n"
        "返回 ``{title, paragraphs, word_count}``。``title`` 来自文档第一段"
        "Heading 1 标题（若无则取首段文本前 40 字）。``paragraphs`` 是除标题"
        "外的所有正文段落。文件不存在返回 ``{error: 'file_not_found'}``。"
    ),
    category="files",
    consumers=[Consumer.CHAT_AGENT],
    trigger_modes=[TriggerMode.ON_DEMAND],
    icon="file-search",
    parameters_schema={
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "要读取的 Word 文件名",
            },
        },
        "required": ["filename"],
    },
)
async def docx_read(filename: str = "", **_kwargs: Any) -> dict:
    if not filename or not filename.strip():
        return {"error": "missing_filename"}
    try:
        canonical = _normalize_filename(filename)
        target = safe_resolve(_safe_dir(), canonical, allow_subdirs=False)
    except ValueError as exc:
        return {"error": "invalid_path", "detail": str(exc)}

    if not target.exists():
        return {"error": "file_not_found", "filename": canonical}

    try:
        doc = Document(str(target))
    except PackageNotFoundError as exc:
        return {"error": "parse_failed", "detail": f"not a valid docx: {exc}"}
    except Exception as exc:
        logger.exception("[docx.read] parse failed for %s", target)
        return {"error": "parse_failed", "detail": str(exc)}

    title: str = ""
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = (para.style.name if para.style else "") or ""
        if not title and style_name.startswith("Heading"):
            title = text
            continue
        paragraphs.append(text)

    if not title and paragraphs:
        # 退化：没用 Heading 时取第一段前 40 字当标题（不要从 paragraphs 移除）
        title = paragraphs[0][:40]

    word_count = sum(len(p) for p in paragraphs)  # 中文按字符数，英文段中各算
    return {
        "title": title,
        "paragraphs": paragraphs,
        "word_count": word_count,
    }


# ---------------------------------------------------------------------------
# 3. docx.append
# ---------------------------------------------------------------------------

@register_capability(
    name="docx.append",
    display_name="向 Word 文档追加内容",
    description=(
        "向已有的 Word 文档末尾追加段落（不破坏原有内容）。适用场景：用户说"
        "「再补一段…」「加上 XX 内容」「补充一句…」。\n\n"
        "参数：\n"
        "- filename: 已存在的 Word 文件名\n"
        "- paragraphs: 要追加的段落列表（list[str]，每段一项）\n\n"
        "返回 ``{path, appended_count, total_paragraphs}``。文件不存在返回 "
        "``{error: 'file_not_found'}`` —— 此时让用户先用 docx.create 创建。"
    ),
    category="files",
    consumers=[Consumer.CHAT_AGENT],
    trigger_modes=[TriggerMode.ON_DEMAND],
    icon="file-plus",
    parameters_schema={
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "已存在的 Word 文件名",
            },
            "paragraphs": {
                "type": "array",
                "items": {"type": "string"},
                "description": "追加的段落列表",
            },
        },
        "required": ["filename", "paragraphs"],
    },
)
async def docx_append(
    filename: str = "",
    paragraphs: list[str] | None = None,
    **_kwargs: Any,
) -> dict:
    if not filename or not filename.strip():
        return {"error": "missing_filename"}
    if not paragraphs:
        return {"error": "empty_paragraphs"}

    try:
        canonical = _normalize_filename(filename)
        target = safe_resolve(_safe_dir(), canonical, allow_subdirs=False)
    except ValueError as exc:
        return {"error": "invalid_path", "detail": str(exc)}

    if not target.exists():
        return {"error": "file_not_found", "filename": canonical}

    try:
        doc = Document(str(target))
    except PackageNotFoundError as exc:
        return {"error": "parse_failed", "detail": f"not a valid docx: {exc}"}
    except Exception as exc:
        logger.exception("[docx.append] parse failed for %s", target)
        return {"error": "parse_failed", "detail": str(exc)}

    appended = 0
    for p in paragraphs:
        text = str(p)
        if not text.strip():
            continue
        doc.add_paragraph(text)
        appended += 1

    try:
        doc.save(str(target))
    except Exception as exc:
        logger.exception("[docx.append] save failed for %s", target)
        return {"error": "save_failed", "detail": str(exc)}

    total = sum(1 for para in doc.paragraphs if (para.text or "").strip())
    return {
        "path": target.name,
        "appended_count": appended,
        "total_paragraphs": total,
    }


__all__ = ["docx_create", "docx_read", "docx_append"]
